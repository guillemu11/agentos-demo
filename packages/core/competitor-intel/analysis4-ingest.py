#!/usr/bin/env python3
"""Parses Analysis 4 .docx, extracts per-heading 'Overall: X.X/10' patterns,
and falls back to table-based extraction (Table 6 format: Brand | vs. EH | Score | Delta).

Usage: python analysis4-ingest.py <path_to_docx>
Output on stdout: JSON array of { brand_name, overall }.
"""
import sys, json, re
from docx import Document

SCORE_RE = re.compile(r"(?:Overall|Score)[^\d]*(\d+\.?\d*)\s*/\s*10", re.I)
# Matches "8/10", "7.5/10", "6 / 10" etc. in a cell
CELL_SCORE_RE = re.compile(r"(\d+\.?\d*)\s*/\s*10")
# Emirates Holidays header pattern: "vs. Emirates Holidays (7.2)"
EH_REF_RE = re.compile(r"Emirates\s+Holidays.*?(\d+\.?\d*)", re.I)


def parse(path):
    doc = Document(path)
    results = []
    seen = set()

    # --- Strategy 1: heading-based paragraph scan ---
    current_brand = None
    buffer = []

    def flush_heading():
        if not current_brand:
            return
        joined = "\n".join(buffer)
        m = SCORE_RE.search(joined)
        if m:
            try:
                name = current_brand
                score = float(m.group(1))
                if name not in seen:
                    seen.add(name)
                    results.append({"brand_name": name, "overall": score})
            except ValueError:
                pass

    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            flush_heading()
            current_brand = txt
            buffer = []
        else:
            buffer.append(txt)
    flush_heading()

    # --- Strategy 2: table scan for scoring table (Brand | vs EH | Overall DX score | ...) ---
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        # detect the scoring table: needs a "Brand" col and an "Overall" col
        brand_col = next((i for i, h in enumerate(header_cells) if re.search(r"brand", h, re.I)), None)
        score_col = next((i for i, h in enumerate(header_cells) if re.search(r"overall|score", h, re.I)), None)
        eh_col = next((i for i, h in enumerate(header_cells) if re.search(r"emirates", h, re.I)), None)

        if brand_col is None or score_col is None:
            continue

        # Extract Emirates Holidays reference from header if present
        if eh_col is not None:
            eh_header = header_cells[eh_col]
            m = EH_REF_RE.search(eh_header)
            if m:
                name = "Emirates Holidays"
                if name not in seen:
                    seen.add(name)
                    results.append({"brand_name": name, "overall": float(m.group(1))})

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(brand_col, score_col):
                continue
            brand = cells[brand_col].strip()
            raw_score = cells[score_col].strip()
            if not brand:
                continue
            m = CELL_SCORE_RE.search(raw_score)
            if m:
                try:
                    score = float(m.group(1))
                    if brand not in seen:
                        seen.add(brand)
                        results.append({"brand_name": brand, "overall": score})
                except ValueError:
                    pass

    # Also check Table 0 for Emirates Holidays headline score
    if doc.tables:
        t0 = doc.tables[0]
        for row in t0.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if re.search(r"emirates\s+holidays", txt, re.I):
                    m = CELL_SCORE_RE.search(txt)
                    if m and "Emirates Holidays" not in seen:
                        seen.add("Emirates Holidays")
                        results.append({"brand_name": "Emirates Holidays", "overall": float(m.group(1))})

    print(json.dumps(results))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("[]")
        sys.exit(0)
    parse(sys.argv[1])
