#!/usr/bin/env python3
"""Generates Analysis 5 .docx from JSON on stdin, writes bytes to stdout."""
import sys, json
from docx import Document
from docx.shared import Pt


def main():
    data = json.loads(sys.stdin.read())
    doc = Document()

    inv = data.get('investigation', {})
    doc.add_heading(f"Analysis 5 — {inv.get('name', 'Lifecycle Audit')}", level=0)
    doc.add_heading("Lifecycle, Email & Customer Journey Assessment", level=1)
    doc.add_paragraph(
        "This analysis extends Analyses 1-4 with first-hand lifecycle evidence: real subscriptions, "
        "real inboxes, real engagement across synthetic personas. Findings are drawn from live "
        "ingestion of emails received, classified by type, and scored across four lifecycle axes."
    )

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    brands = data.get('brands', [])
    scored = [b for b in brands if (b.get('scores') or {}).get('overall') is not None]
    if scored:
        avg = sum(b['scores']['overall'] for b in scored) / len(scored)
        doc.add_paragraph(
            f"Portfolio lifecycle average: {avg:.1f}/10 across {len(scored)} scored brands "
            f"({len(brands)} investigated)."
        )
    else:
        doc.add_paragraph(
            f"{len(brands)} brands investigated. Scoring in progress."
        )

    # Per-brand
    for b in brands:
        doc.add_heading(b.get('name', 'Unnamed'), level=1)
        if b.get('positioning'):
            p = doc.add_paragraph()
            run = p.add_run(b['positioning'])
            run.italic = True
        s = b.get('scores') or {}

        def _fmt(v):
            return f"{v:.1f}" if isinstance(v, (int, float)) else "—"

        doc.add_paragraph(
            f"Lifecycle maturity: {_fmt(s.get('lifecycle_maturity'))}/10   ·   "
            f"Email sophistication: {_fmt(s.get('email_sophistication'))}/10   ·   "
            f"Journey depth: {_fmt(s.get('journey_depth'))}/10   ·   "
            f"Personalisation: {_fmt(s.get('personalisation'))}/10   ·   "
            f"Overall: {_fmt(s.get('overall'))}/10"
        )

        if s.get('manual_notes'):
            doc.add_paragraph(s['manual_notes'])

        insights = b.get('insights') or []
        if insights:
            for ins in insights:
                doc.add_heading(ins['title'], level=2)
                if ins.get('body'):
                    doc.add_paragraph(ins['body'])

    # Cross-brand insights
    cross = data.get('cross_brand_insights', [])
    if cross:
        doc.add_heading("Cross-brand insights", level=1)
        for ins in cross:
            doc.add_heading(ins['title'], level=2)
            if ins.get('body'):
                doc.add_paragraph(ins['body'])

    # Comparative — Time to first useful email
    comp = data.get('comparative', {})
    ttft = comp.get('ttft', [])
    if ttft:
        doc.add_heading("Comparative — Time to first useful email", level=1)
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Light Grid Accent 1'
        t.rows[0].cells[0].text = 'Brand'
        t.rows[0].cells[1].text = 'Time to first useful email'
        for r in ttft:
            row = t.add_row().cells
            row[0].text = r.get('brand_name', '—')
            sec = r.get('seconds_to_first')
            if sec is None:
                row[1].text = 'No email received'
            elif sec < 60:
                row[1].text = f"{int(sec)}s"
            elif sec < 3600:
                row[1].text = f"{int(sec/60)} min"
            elif sec < 86400:
                row[1].text = f"{sec/3600:.1f} h"
            else:
                row[1].text = f"{sec/86400:.1f} d"

    # Quick wins
    qw = data.get('quick_wins', [])
    if qw:
        doc.add_heading("Quick wins", level=1)
        for w in qw:
            doc.add_paragraph(f"• {w}")

    # Methodology note
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Two synthetic personas were created and authenticated via Gmail. "
        "Newsletter subscriptions, preference-center interactions and abandonment tests were "
        "executed manually on each brand's web properties. Emails were ingested via Gmail API "
        "and classified in two phases: deterministic domain matching and Claude Sonnet LLM "
        "fallback for ambiguous cases. Engagement (opens, clicks) was simulated per a persona-specific "
        "pattern to trigger downstream journeys where available. Scoring combines automated email-based "
        "heuristics with manual analyst review."
    )

    doc.save(sys.stdout.buffer)


if __name__ == "__main__":
    main()
